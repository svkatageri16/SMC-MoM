import os, time, smtplib, webbrowser, uvicorn, re, gc, json
from dotenv import load_dotenv
from datetime import datetime

# Load environment variables
load_dotenv()
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pyngrok import ngrok
from email.message import EmailMessage
from faster_whisper import WhisperModel
import google.generativeai as genai
from docx import Document

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# --- CONFIGURATION ---
BASE_DIR = r"D:\SMC MoM\sessions"
os.makedirs(BASE_DIR, exist_ok=True)
METADATA_FILE = os.path.join(BASE_DIR, "sessions.json")

NGROK_AUTH_TOKEN = os.getenv("NGROK_AUTH_TOKEN", "#")
EMAIL_USER = os.getenv("EMAIL_USER", "#@gmail.com")
EMAIL_PASS = os.getenv("EMAIL_PASS", "#") 
MLA_EMAILS = os.getenv("MLA_EMAILS", "#").split(",")

# Persistent State (Will be updated by Frontend session_id)
state = {"current_session": {"id": None, "path": None, "title": None}}
admin_connections = []
public_url = None

# --- AI ENGINE ---
print("--- 🧠 Initializing GPU AI (Large-V3) ---")
gc.collect()
try:
    whisper_model = WhisperModel("large-v3", device="cuda", compute_type="float16")
except:
    whisper_model = WhisperModel("medium", device="cpu", compute_type="int8")

genai.configure(api_key=os.getenv("GEMINI_API_KEY", "#"))

app.mount("/data", StaticFiles(directory=BASE_DIR), name="data")

def get_all_sessions():
    if not os.path.exists(METADATA_FILE): return []
    try:
        with open(METADATA_FILE, "r") as f: 
            sessions = json.load(f)
        return [s for s in sessions if os.path.exists(s.get("path", ""))]
    except: return []

def save_session_metadata(session_data):
    sessions = get_all_sessions()
    exists = False
    for i, s in enumerate(sessions):
        if s["id"] == session_data["id"]:
            sessions[i] = session_data
            exists = True
            break
    if not exists: sessions.append(session_data)
    with open(METADATA_FILE, "w", encoding="utf-8") as f: json.dump(sessions, f, indent=4)

@app.on_event("startup")
async def startup_event():
    global public_url
    try:
        ngrok.set_auth_token(NGROK_AUTH_TOKEN)
        public_url = ngrok.connect(8000).public_url
        print(f"🚀 Digital Sansad Ready. Public Link: {public_url}")
    except: print("Ngrok connection failed.")

@app.get("/api/sessions")
async def list_sessions():
    return sorted(get_all_sessions(), key=lambda x: x['id'], reverse=True)

@app.get("/api/get_transcript")
async def get_saved_transcript(session_id: str):
    path = os.path.join(BASE_DIR, session_id, "transcript.txt")
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return {"transcript": f.read()}
    return {"transcript": ""}

@app.post("/api/sessions/start")
async def start_new_session(request: Request):
    global public_url
    data = await request.json()
    title = data.get("title", "Untitled_Meeting")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    clean_title = re.sub(r'\W+', '', title)
    session_id = f"{ts}_{clean_title}"
    session_path = os.path.join(BASE_DIR, session_id)
    os.makedirs(os.path.join(session_path, "recordings"), exist_ok=True)
    
    state["current_session"] = {
        "id": session_id, "path": session_path, "title": title, 
        "date": datetime.now().strftime("%d %b %Y, %H:%M"),
        "mom_ready": False, "transcript_ready": False
    }
    save_session_metadata(state["current_session"])

    # TRIGGER EMAILS
    try:
        msg = EmailMessage()
        msg.set_content(f"Digital Sansad Official Session Started: {title}\nSpeaker Link: {public_url}")
        msg['Subject'] = f'LIVE SESSION: {title}'
        msg['From'] = EMAIL_USER
        msg['To'] = ", ".join(MLA_EMAILS)
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(EMAIL_USER, EMAIL_PASS)
            smtp.send_message(msg)
        print("✅ Emails sent to members.")
    except Exception as e: print(f"Email Dispatch Error: {e}")

    return state["current_session"]

@app.get("/api/files")
async def get_files(session_id: str):
    rec_path = os.path.join(BASE_DIR, session_id, "recordings")
    files_list = []
    if os.path.exists(rec_path):
        for f in os.listdir(rec_path):
            if f.endswith(".webm"):
                parts = f.replace(".webm", "").split("_")
                if len(parts) >= 3:
                    files_list.append({
                        "name": parts[0], "designation": parts[1], 
                        "timestamp": parts[2], "filename": f, "sid": session_id
                    })
    return sorted(files_list, key=lambda x: int(x['timestamp']), reverse=True)

@app.get("/api/transcribe")
async def transcribe_all(session_id: str):
    session_path = os.path.join(BASE_DIR, session_id)
    rec_path = os.path.join(session_path, "recordings")
    files = sorted([f for f in os.listdir(rec_path) if f.endswith(".webm")], 
                   key=lambda x: int(x.replace(".webm", "").split("_")[2]))
    
    final_transcript = ""
    # MULTI-LANGUAGE PROMPT (Anchor)
    prompt = ("नमस्कार, सोलापूर महानगरपालिका बैठकीत आपले स्वागत आहे. आज आपण विकास कामांवर चर्चा करणार आहोत. Please speak in Marathi, Hindi or English."
    "नमस्कार, आजच्या बैठकीत आपले स्वागत आहे. "  # Marathi opening
    "सब लोग कृपया ध्यान दें, आज हम इस विषय पर चर्चा करेंगे. "  # Hindi transition
    "Please ensure all points are noted in English. "  # English transition
    "आपल्याला हे काम लवकरात लवकर पूर्ण करायचे आहे, ताकि परिणाम अच्छे मिलें. " # Mixed switch
    "This is an official SMC meeting record in Marathi, Hindi, and English." # Summary
    )

    for f in files:
        file_path = os.path.join(rec_path, f)
        parts = f.replace(".webm", "").split("_")
        speaker_name = parts[0].replace(".", " ")
        desig = parts[1]
        readable_time = datetime.fromtimestamp(int(parts[2])).strftime('%H:%M:%S')
        
        try:
            print(f"🚀 GPU Transcribing: {speaker_name} ({f})")
            segments, _ = whisper_model.transcribe(
                file_path, initial_prompt=prompt, 
                beam_size=5, temperature=0, language=None, vad_filter=True,
                condition_on_previous_text=False # Prevents repeating the prompt
            )
            text = " ".join([s.text.strip() for s in segments])
            if text.strip():
                final_transcript += f"[{speaker_name} ({desig}), {readable_time}]: {text.strip()}\n\n"
        except Exception as e:
            print(f"⚠️ Skipping corrupted file {f}: {e}")
            continue
    
    if final_transcript:
        with open(os.path.join(session_path, "transcript.txt"), "w", encoding="utf-8") as tfile:
            tfile.write(final_transcript)
        
        sessions = get_all_sessions()
        for s in sessions:
            if s["id"] == session_id: s["transcript_ready"] = True
        with open(METADATA_FILE, "w", encoding="utf-8") as f: json.dump(sessions, f, indent=4)
        
    return {"transcript": final_transcript}

@app.post("/api/gemini")
async def generate_mom(request: Request):
    data = await request.json()
    transcript = data.get("transcript")
    sid = data.get("session_id")
    session_data = next((s for s in get_all_sessions() if s["id"] == sid), None)
    
    prompt = f"""
    ROLE: Professional Parliamentary Scribe for Solapur Municipal Corporation.
    TASK: Generate a formal MoM from the PROVIDED TRANSCRIPT.
    STRICT RULES:
    1. ZERO HALLUCINATION. Do not add speakers, attendees, or topics not mentioned in the transcript.
    2. DATE: {session_data['date']}.
    3. LANGUAGE: Strictly English.
    4. FORMAT: Header, Roll Call (only who spoke), Summary, Resolution Table.
    
    TRANSCRIPT:
    {transcript}
    """
    
    try:
        llm = genai.GenerativeModel('gemini-2.5-flash')
        response = llm.generate_content(prompt)
        content = response.text
    except Exception as e:
        print(f"Gemini Fail: {e}")
        content = f"OFFICIAL SESSION RECORD (OFFLINE VERIFIED)\nDate: {session_data['date']}\n\n{transcript}"

    filename = f"MoM_{sid}.docx"
    save_path = os.path.join(BASE_DIR, sid, filename)
    doc = Document()
    doc.add_heading(f"SMC OFFICIAL RECORD - {session_data['title']}", 0)
    doc.add_paragraph(f"Date: {session_data['date']}")
    doc.add_paragraph(content)
    doc.save(save_path)
    
    sessions = get_all_sessions()
    for s in sessions:
        if s["id"] == sid: s["mom_ready"] = True
    with open(METADATA_FILE, "w", encoding="utf-8") as f: json.dump(sessions, f, indent=4)
    
    os.startfile(save_path)
    return {"status": "success"}

@app.get("/api/open_mom")
async def open_mom(session_id: str):
    path = os.path.join(BASE_DIR, session_id, f"MoM_{session_id}.docx")
    if os.path.exists(path):
        os.startfile(path)
        return {"status": "success"}
    return JSONResponse({"error": "File Not Found"}, status_code=404)

@app.websocket("/ws/stream")
async def stream_audio(websocket: WebSocket, name: str = "Unknown", desig: str = "Officer"):
    await websocket.accept()
    curr = state["current_session"]
    if not curr or not curr["id"]:
        await websocket.close(); return

    ts = int(time.time())
    safe_name = name.replace(" ", ".")
    filename = f"{safe_name}_{desig}_{ts}.webm"
    filepath = os.path.join(curr["path"], "recordings", filename)
    
    await broadcast_to_admins({"event": "NEW_SPEAKER", "name": name})
    with open(filepath, "wb") as f:
        try:
            while True:
                data = await websocket.receive_bytes(); f.write(data)
        except WebSocketDisconnect:
            await broadcast_to_admins({"event": "STREAM_FINISHED", "filename": filename})

@app.websocket("/ws/admin")
async def admin_socket(websocket: WebSocket):
    await websocket.accept(); admin_connections.append(websocket)
    try:
        while True: await websocket.receive_text()
    except WebSocketDisconnect: admin_connections.remove(websocket)

async def broadcast_to_admins(message: dict):
    for conn in admin_connections: 
        try: await conn.send_json(message)
        except: admin_connections.remove(conn)

@app.get("/") 
async def get_speaker():
    with open(r"D:\SMC MoM\index.html", "r", encoding="utf-8") as f: return HTMLResponse(f.read())

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)