from faster_whisper import WhisperModel
import google.generativeai as genai
from docx import Document

# Initialize AI
whisper_model = WhisperModel("medium", device="cpu", compute_type="int8")

@app.get("/api/transcribe")
async def transcribe_all():
    files = sorted([f for f in os.listdir(RECORDING_DIR) if f.endswith(".webm")], 
                   key=lambda x: int(x.split("_")[2]))
    full_transcript = ""
    for f in files:
        name = f.split("_")[0]
        segments, _ = whisper_model.transcribe(
            os.path.join(RECORDING_DIR, f), 
            initial_prompt="This is a Parliamentary session in Hindi, Marathi, and English."
        )
        for s in segments:
            full_transcript += f"[{name}]: {s.text}\n"
    return {"transcript": full_transcript}

@app.post("/api/generate-mom")
async def generate_mom(data: dict):
    transcript = data.get("transcript")
    # Gemini Logic
    llm = genai.GenerativeModel('models/gemini-3-flash-preview')
    response = llm.generate_content(f"Create formal Parliamentary MoM for: {transcript}")
    
    # Docx Generation
    doc = Document()
    doc.add_heading('Official Meeting Minutes', 0)
    doc.add_paragraph(response.text)
    save_path = r"D:\SMC MoM\Final_MoM.docx"
    doc.save(save_path)
    os.startfile(save_path) # Auto-opens Word
    return {"status": "success"}