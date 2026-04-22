import os
import re
import google.generativeai as genai
from faster_whisper import WhisperModel
from docx import Document

# --- 1. CONFIGURATION ---
GEMINI_KEY = "AIzaSyBBpZkIdcy7Il9arcoKWaoxnRSWecIeuR8" # Ensure this is your correct key
genai.configure(api_key=GEMINI_KEY)

def get_file_time(filename):
    match = re.search(r"_(\d+)\.", filename)
    if match:
        return int(match.group(1))
    return os.path.getctime(os.path.join(r"D:\SMC MoM\recordings", filename))

def run_it():
    rec_path = r"D:\SMC MoM\recordings"
    output_docx = r"D:\SMC MoM\Final_MoM.docx"
    
    files = [f for f in os.listdir(rec_path) if f.endswith(".webm")]
    files.sort(key=get_file_time)
    
    if not files:
        print("❌ No recordings found.")
        return

    print(f"--- 📂 Processing {len(files)} clips ---")

    # LOAD MODEL
    print("--- 🧠 Loading AI Engine (CPU Mode) ---")
    model = WhisperModel("medium", device="cpu", compute_type="int8")

    full_transcript = []
    solapur_context = "Meeting in Solapur. Languages: Marathi, Hindi, Kannada, English. मराठी, हिंदी, ಕನ್ನಡ."

    for f in files:
        speaker = f.split('_')[0]
        file_path = os.path.join(rec_path, f)
        print(f"🎙️  Transcribing {speaker}...")
        
        try:
            segments, info = model.transcribe(
                file_path, 
                beam_size=5,
                initial_prompt=solapur_context
            )
            for s in segments:
                line = f"[{speaker}]: {s.text.strip()}"
                print(line)
                full_transcript.append(line)
        except Exception as e:
            print(f"⚠️ Error on {f}: {e}")

    if not full_transcript:
        print("⚠️ Transcript empty.")
        return

    # --- THE CRITICAL GEMINI FIX ---
    print("\n--- ✍️  Generating Official MoM via Gemini ---")
    transcript_text = "\n".join(full_transcript)
    
    # Save the text as a backup FIRST
    with open(r"D:\SMC MoM\raw_transcript.txt", "w", encoding="utf-8") as f:
        f.write(transcript_text)

    try:
        # FIX: Try the direct 'gemini-1.5-flash' name
        # If this fails, the system will tell us which models ARE available
        llm = genai.GenerativeModel('models/gemini-3-flash-preview')
        
        prompt = f"""
        Act as a professional Parliamentary Scribe.
        Format this multilingual transcript (Marathi/Hindi/Kannada/English) into official MoM.
        
        TRANSCRIPT:
        {transcript_text}
        """
        
        response = llm.generate_content(prompt)
        
        # Word Doc Logic
        doc = Document()
        doc.add_heading('SMC - OFFICIAL PROCEEDINGS', 0)
        doc.add_paragraph(response.text)
        doc.save(output_docx)
        print(f"✅ SUCCESS! File created: {output_docx}")
        
    except Exception as e:
        print(f"❌ Gemini API Error: {e}")
        print("Checking available models for your API key...")
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                print(f"Available Model: {m.name}")

if __name__ == "__main__":
    run_it()