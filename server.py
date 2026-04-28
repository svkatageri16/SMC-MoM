import os, time, re, uvicorn, json
from dotenv import load_dotenv
from fastapi import FastAPI, WebSocket, Request

# Load environment variables
load_dotenv()
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from faster_whisper import WhisperModel
import google.generativeai as genai
from docx import Document

app = FastAPI()

# PATHS
BASE_DIR = r"D:\SMC MoM"
REC_DIR = os.path.join(BASE_DIR, "recordings")
os.makedirs(REC_DIR, exist_ok=True)

# Mount recordings so the dashboard can play them
app.mount("/recordings", StaticFiles(directory=REC_DIR), name="recordings")

# AI CONFIG
GEMINI_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_KEY)
# Initialize Whisper once on startup to save time
print("--- Initializing Whisper (CPU/Medium) ---")
whisper_model = WhisperModel("medium", device="cpu", compute_type="int8")

@app.get("/") # Speaker Page
async def get_speaker():
    with open("index.html", "r", encoding="utf-8") as f: return HTMLResponse(f.read())

@app.get("/master") # Admin Dashboard
async def get_admin():
    with open("admin.html", "r", encoding="utf-8") as f: return HTMLResponse(f.read())

@app.websocket("/ws/{name}/{desig}")
async def websocket_endpoint(websocket: WebSocket, name: str, desig: str):
    await websocket.accept()
    ts = int(time.time())
    # Filename structure: Name_Designation_Timestamp.webm
    filename = f"{name}_{desig}_{ts}.webm"
    file_path = os.path.join(REC_DIR, filename)
    with open(file_path, "wb") as f:
        try:
            while True:
                data = await websocket.receive_bytes()
                f.write(data)
        except: print(f"Finished: {filename}")

@app.get("/api/files")
async def list_files():
    files = []
    for f in sorted(os.listdir(REC_DIR)):
        if f.endswith(".webm"):
            parts = f.replace(".webm", "").split("_")
            if len(parts) >= 3:
                files.append({
                    "name": parts[0], "designation": parts[1], 
                    "time": int(parts[2]), "filename": f
                })
    return JSONResponse(files)

@app.get("/api/process")
async def process_recordings():
    transcript = ""
    files = sorted([f for f in os.listdir(REC_DIR) if f.endswith(".webm")], 
                   key=lambda x: int(x.split("_")[2]))
    
    for f in files:
        parts = f.split("_")
        name, desig = parts[0], parts[1]
        segments, _ = whisper_model.transcribe(
            os.path.join(REC_DIR, f), language="hi",
            initial_prompt="Meeting in Solapur (Marathi, Hindi, English). No Urdu."
        )
        for s in segments:
            transcript += f"[{name} - {desig}]: {s.text}\n"
    
    return {"transcript": transcript}

@app.post("/api/gemini")
async def call_gemini(request: Request):
    data = await request.json()
    transcript = data.get("transcript")
    
    # Using the specific model you requested
    llm = genai.GenerativeModel('models/gemini-3-flash-preview')
    prompt = f"Create a professional Parliamentary MoM from this edited transcript:\n\n{transcript}"
    response = llm.generate_content(prompt)
    
    # Save to Word
    doc = Document()
    doc.add_heading('SMC OFFICIAL PROCEEDINGS', 0)
    doc.add_paragraph(response.text)
    doc_path = os.path.join(BASE_DIR, "Final_MoM.docx")
    doc.save(doc_path)
    
    # AUTOMATICALLY OPEN THE FILE (Windows only)
    os.startfile(doc_path)
    
    return {"path": doc_path}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)